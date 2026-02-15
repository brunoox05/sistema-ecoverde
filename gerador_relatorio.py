from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import sqlite3
import math

# --- ESTILOS ---
def configurar_estilo_global(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri Light'
    font.size = Pt(12)
    p_format = style.paragraph_format
    p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_format.line_spacing = 1.5
    p_format.space_after = Pt(6)

def formatar_titulo(paragrafo):
    for run in paragrafo.runs:
        run.font.name = 'Calibri Light'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

def colorir_celula(cell, hex_color):
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hex_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def formatar_texto_tabela(cell, texto, negrito=False, cor_fonte=None, alinhamento='center', tamanho=11):
    cell.text = str(texto)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if alinhamento == 'center' else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = 'Calibri Light'
            run.font.size = Pt(tamanho)
            run.font.bold = negrito
            if cor_fonte: run.font.color.rgb = cor_fonte

def calcular_tempo_caminhada(distancia):
    if distancia == 0: return "0 minutos"
    return f"{math.ceil(distancia / 80)} minutos"

# --- GERADOR ---
def gerar_docx(empresa_id, nome_empresa):
    conn = sqlite3.connect('dados_ric_oficial.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    doc = Document()
    configurar_estilo_global(doc)
    
    # Cores
    COR_PRINCIPAL = "548235"
    COR_LINHA_CLARA = "E2EFDA"
    COR_FONTE_BRANCA = RGBColor(255, 255, 255)
    COR_FONTE_PRETA = RGBColor(0, 0, 0)
    COR_TITULO_VERMELHO = RGBColor(255, 0, 0)

    # Título
    titulo = doc.add_heading(f'RELATÓRIO DE IMPACTO NA CIRCULAÇÃO - {nome_empresa.upper()}', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs: 
        run.font.name = 'Calibri Light'; run.font.bold = True; run.font.color.rgb = RGBColor(0,0,0)

    # 1. DESCRIÇÃO PEDS
    h1 = doc.add_heading('1. Descrição dos PEDs', level=1); formatar_titulo(h1)
    
    cursor.execute(f"SELECT * FROM peds WHERE empresa_id = {empresa_id} ORDER BY numero_ped ASC")
    peds = cursor.fetchall()

    if not peds: doc.add_paragraph("Nenhum PED cadastrado.")

    for ped in peds:
        cursor.execute(f"SELECT r.* FROM ruas r JOIN ped_ruas pr ON r.id = pr.rua_id WHERE pr.ped_id = {ped['id']}")
        ruas_ped = cursor.fetchall()
        
        txt_abrigo = "sem abrigo" if ped['tem_abrigo'] == "Não" else f"com abrigo ({ped['tipo_abrigo'].lower()})"
        txt_assento = "sem assento" if ped['tem_assento'] == "Não" else f"com assento"
        
        tem_vert = any("Existente" in (r['sinalizacao_vertical'] or "") for r in ruas_ped)
        tem_horiz = any("Existente" in (r['sinalizacao_horizontal'] or "") for r in ruas_ped)
        
        txt_sinal_v = "Possui sinalização vertical" if tem_vert else "Não possui sinalização vertical"
        txt_sinal_h = "apresenta sinalização horizontal" if tem_horiz else "não apresenta sinalização horizontal"

        lista_p = []
        estado_geral = "boas condições"
        for r in ruas_ped:
            if r['problemas_calcada']: lista_p.extend([p.strip().lower() for p in r['problemas_calcada'].split(",")])
            if r['estado_calcada'] in ['Ruim', 'Péssimo']: estado_geral = "condições ruins"
        
        lista_p = list(set(lista_p))
        if lista_p:
            if len(lista_p) > 1: probs_txt = ", ".join(lista_p[:-1]) + " e " + lista_p[-1]
            else: probs_txt = lista_p[0]
            txt_calcada = f"A calçada encontra-se em {estado_geral}, evidenciada pela presença de {probs_txt}."
        else:
            txt_calcada = "A calçada encontra-se em boas condições."

        p = doc.add_paragraph()
        dist_int = int(ped['distancia_empreendimento'])
        run = p.add_run(f"O PED-{ped['numero_ped']:02d} é apresentado {txt_abrigo} e {txt_assento}. {txt_sinal_v}, {txt_sinal_h}. {txt_calcada} Destaca-se que o PED está localizado a uma distância de {dist_int} metros do empreendimento.")
        run.font.name = 'Calibri Light'

    # 2. CAMINHAMENTO
    doc.add_page_break()
    h2 = doc.add_heading('2. Análise do Caminhamento', level=1); formatar_titulo(h2)
    
    for ped in peds:
        p_sub = doc.add_paragraph()
        run_sub = p_sub.add_run(f"Análise do Caminhamento - PED {ped['numero_ped']:02d}")
        run_sub.bold = True; run_sub.font.name = 'Calibri Light'; run_sub.font.color.rgb = COR_TITULO_VERMELHO
        p_sub.paragraph_format.space_after = Pt(0)
        
        cursor.execute(f"SELECT r.* FROM ruas r JOIN ped_ruas pr ON r.id = pr.rua_id WHERE pr.ped_id = {ped['id']}")
        ruas_ped = cursor.fetchall()

        # Pavimentação
        p_pav = doc.add_paragraph()
        run_bold = p_pav.add_run("Pavimentação: ")
        run_bold.bold = True; run_bold.font.name = 'Calibri Light'
        
        for i, r in enumerate(ruas_ped):
            prefixo = " " if i > 0 else ""
            
            # CORREÇÃO: Remoção do .get() que causava o erro
            # O sistema agora confia que a coluna existe no banco
            if r['tem_pavimento'] == "Não":
                texto = f"Na via {r['nome']}, observa-se que não possui pavimentação, apresenta piso de {r['tipo_pavimento'].lower()}. "
            else:
                texto = f"Na via {r['nome']}, observa-se que é devidamente pavimentada ({r['tipo_pavimento'].lower()}). "
            
            if r['problemas_pavimento']: 
                texto += f"A pavimentação apresenta obstruções durante o percurso, neste caso {r['problemas_pavimento'].lower()}. "
            else: 
                texto += "A pavimentação não apresenta obstruções. "
            
            if "Inexistente" in r['sinalizacao_horizontal']:
                texto += "Não possui sinalização horizontal na via."
            else:
                cond_h = "más condições" if "Ruim" in r['sinalizacao_horizontal'] else "boas condições"
                texto += f"Possui sinalização horizontal na via em {cond_h}."
            
            p_pav.add_run(prefixo + texto)

        # Passeio
        p_pass = doc.add_paragraph()
        run_bold = p_pass.add_run("Condições do Passeio: ")
        run_bold.bold = True; run_bold.font.name = 'Calibri Light'
        
        for i, r in enumerate(ruas_ped):
            prefixo = " " if i > 0 else ""
            texto = f"Na {r['nome']}, "
            
            if r['tem_calcada'] == "Inexistente": texto += "não há calçamento. "
            else:
                texto += "há calçamento. "
                if r['problemas_calcada']: texto += f"O calçamento apresenta obstruções causadas por {r['problemas_calcada'].lower()}. "
                else: texto += "O calçamento não apresenta obstruções. "
                
                if r['tem_acessibilidade'] == "Sim": 
                    acess_desc = r['tipos_acessibilidade'].lower() if r['tipos_acessibilidade'] else "itens de acessibilidade"
                    texto += f"Há adequações de acessibilidade, neste caso possui {acess_desc}. "
                else: texto += "Não há adequações de acessibilidade. "
                
            texto += "Não possui sinalização vertical." if "Inexistente" in r['sinalizacao_vertical'] else "Possui sinalização vertical na via e está em boas condições."
            p_pass.add_run(prefixo + texto)

        # Avaliação
        p_aval = doc.add_paragraph()
        run_bold = p_aval.add_run("Avaliação do percurso: ")
        run_bold.bold = True; run_bold.font.name = 'Calibri Light'
        
        dist_int = int(ped['distancia_empreendimento'])
        tempo = calcular_tempo_caminhada(dist_int)
        
        p_aval.add_run(f"Durante o deslocamento da empresa até o ponto PED {ped['numero_ped']:02d}, o pedestre percorre aproximadamente {dist_int} metros em um tempo estimado de {tempo}. De acordo com os parâmetros da Tabela 9, esse percurso é classificado como ")
        run_c = p_aval.add_run(ped['classificacao_tabela'].split("(")[0].strip())
        run_c.bold = True; run_c.font.name = 'Calibri Light'
        p_aval.add_run(".")

    # 3. QUADRO RESUMO
    doc.add_page_break()
    h3 = doc.add_heading('3. Quadro Resumo', level=1); formatar_titulo(h3)
    
    if peds:
        table = doc.add_table(rows=12, cols=len(peds)+1); table.style = 'Table Grid'
        
        colorir_celula(table.cell(0,0), COR_PRINCIPAL)
        labels = {1:"Distância até o empreendimento e\nAvaliação do Caminhamento", 3:"Condições do Abrigo", 4:"Condições dos Assentos", 5:"Condições da Sinalização Vertical", 6:"Condições da Sinalização Horizontal", 7:"Condições da Calçada", 10:"Largura da calçada", 11:"Linhas que atendem"}
        for r_idx, txt in labels.items():
            cell = table.cell(r_idx, 0)
            formatar_texto_tabela(cell, txt, True, COR_FONTE_BRANCA)
            colorir_celula(cell, COR_PRINCIPAL)
            if r_idx == 1: cell.merge(table.cell(2,0))
            if r_idx == 7: cell.merge(table.cell(9,0))
        
        for i, ped in enumerate(peds):
            col = i + 1
            formatar_texto_tabela(table.cell(0, col), f"PED {ped['numero_ped']:02d}", True, COR_FONTE_BRANCA)
            colorir_celula(table.cell(0, col), COR_PRINCIPAL)
            
            cursor.execute(f"SELECT * FROM ruas r JOIN ped_ruas pr ON r.id = pr.rua_id WHERE pr.ped_id = {ped['id']}")
            ruas = cursor.fetchall()
            
            dist_int = int(ped['distancia_empreendimento'])
            txt_class = ped['classificacao_tabela'].split("(")[0].strip()
            
            if ped['tem_abrigo'] == "Não": txt_abrigo = "Não possui abrigo"
            else: txt_abrigo = f"{ped['tipo_abrigo']} em {ped['condicao_abrigo'].lower()} condições"
            
            if ped['tem_assento'] == "Não": txt_assento = "Não possui assento"
            else: txt_assento = f"{ped['tipo_assento']} em {ped['condicao_assento'].lower()} condições"
            
            sin_v = "Possui sinalização vertical em bom estado" if any("Boa" in (r['sinalizacao_vertical'] or "") for r in ruas) else ("Possui sinalização vertical em mau estado" if any("Ruim" in (r['sinalizacao_vertical'] or "") for r in ruas) else "Não possui sinalização vertical")
            sin_h = "Possui sinalização horizontal em bom estado" if any("Boa" in (r['sinalizacao_horizontal'] or "") for r in ruas) else ("Possui sinalização horizontal em mau estado" if any("Ruim" in (r['sinalizacao_horizontal'] or "") for r in ruas) else "Não possui sinalização horizontal")
            
            larguras = [r['largura_calcada'] for r in ruas if r['largura_calcada']]
            media = sum(larguras)/len(larguras) if larguras else 0
            txt_tipo_c = "Calçada Estreita" if media < 1.2 else "Calçada Ampla"
            txt_acess = "Possui acessibilidade" if any(r['tem_acessibilidade']=="Sim" for r in ruas) else "Não possui acessibilidade"
            cond_c = "Em condições ruins" if any(r['estado_calcada'] in ['Ruim','Péssimo'] for r in ruas) else ("Em condições aceitáveis" if any(r['estado_calcada'] == 'Regular' for r in ruas) else "Em boas condições")
            
            dados_col = {1: f"{dist_int}m", 2: txt_class, 3: txt_abrigo, 4: txt_assento, 5: sin_v, 6: sin_h, 7: txt_tipo_c, 8: txt_acess, 9: cond_c, 10: f"{media:.1f}m".replace(".", ","), 11: str(ped['linhas_onibus'])}
            
            for r_idx, val in dados_col.items():
                cell = table.cell(r_idx, col)
                formatar_texto_tabela(cell, val, False, COR_FONTE_PRETA)
                colorir_celula(cell, COR_LINHA_CLARA if r_idx % 2 != 0 else "FFFFFF")

    # 4. CLASSIFICAÇÃO
    doc.add_page_break()
    cursor.execute(f"SELECT r.* FROM ruas r JOIN empresa_ruas er ON r.id = er.rua_id WHERE er.empresa_id = {empresa_id}")
    ruas_todas = cursor.fetchall()
    
    h4 = doc.add_heading('4. Classificação das Vias', level=1); formatar_titulo(h4)
    if ruas_todas:
        t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
        formatar_texto_tabela(t.rows[0].cells[0], "Nome da Via", True, COR_FONTE_BRANCA); colorir_celula(t.rows[0].cells[0], COR_PRINCIPAL)
        formatar_texto_tabela(t.rows[0].cells[1], "Classificação Viária", True, COR_FONTE_BRANCA); colorir_celula(t.rows[0].cells[1], COR_PRINCIPAL)
        for r in ruas_todas:
            row = t.add_row().cells
            formatar_texto_tabela(row[0], r['nome'], alinhamento='left')
            formatar_texto_tabela(row[1], r['classificacao_viaria'])
    
    doc.add_paragraph()
    
    # 5. DETALHAMENTO
    h5 = doc.add_heading('5. Detalhamento das Vias', level=1); formatar_titulo(h5)
    for r in ruas_todas:
        p_t = doc.add_paragraph(f"Ficha Técnica: {r['nome']}"); p_t.runs[0].bold = True; p_t.runs[0].font.name = 'Calibri Light'
        t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'
        
        # CORREÇÃO DO TEXTO DE PAVIMENTAÇÃO
        if r['tem_pavimento'] == "Não":
            pav_txt = f"Não possui pavimentação, apresenta piso de {r['tipo_pavimento'].lower()}."
        else:
            cond = "más" if r['estado_pavimento'] in ["Ruim", "Péssimo"] else "boas"
            pav_txt = f"Pavimentação ({r['tipo_pavimento'].lower()}) em {cond} condições."
            if r['problemas_pavimento']: pav_txt += f" Apresenta {r['problemas_pavimento'].lower()}."
            else: pav_txt += " Não apresenta buracos."

        sin_h_txt = f"Sinalização horizontal {r['sinalizacao_horizontal'].lower()}."
        if "Boa" in r['sinalizacao_horizontal']: sin_h_txt += " Pinturas visíveis e conservadas."
        elif "Ruim" in r['sinalizacao_horizontal']: sin_h_txt += " Pinturas apagadas."
        
        sin_v_txt = f"Sinalização vertical {r['sinalizacao_vertical'].lower()}."
        if "Boa" in r['sinalizacao_vertical']: sin_v_txt += " Placas visíveis e íntegras."
        
        faixa = r['tem_faixa_estacionamento'] if r['tem_faixa_estacionamento'] != "Não" else "Não possui"
        comp = r['local_estacionamento'] if r['estacionamento_irregular'] == "Sim" else "Sem estacionamento irregular"

        dados = [
            ("Logradouro", r['nome']), ("Classificação", r['classificacao_viaria']),
            ("Dimensão", f"{r['largura_via']}m"), ("Sentido", r['sentido_via']),
            ("Faixa Estac.", faixa), ("Faixas Rol.", str(r['num_faixas'])),
            ("Comportamento dos Usuários", comp), ("Pavimentação", pav_txt),
            ("Sinalização Horizontal", sin_h_txt), ("Sinalização Vertical", sin_v_txt)
        ]
        
        for k, v in dados:
            row = t.add_row().cells
            formatar_texto_tabela(row[0], k, True, alinhamento='left')
            formatar_texto_tabela(row[1], v, alinhamento='left')
        doc.add_paragraph()

    conn.close()
    nome_arq = f"Relatorio_RIC_{nome_empresa}.docx"
    doc.save(nome_arq)
    return nome_arq